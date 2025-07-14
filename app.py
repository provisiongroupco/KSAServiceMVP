import streamlit as st
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
import os
from PIL import Image
from streamlit_drawable_canvas import st_canvas
import numpy as np
from utils import (add_header_with_logo, add_footer, style_heading, 
                  create_info_table, add_logo_to_doc, set_cell_margins)
from equipment_config import EQUIPMENT_TYPES

# Page configuration
st.set_page_config(
    page_title="Halton KSA Service Reports",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
    <style>
    .main-header {
        font-size: 2.2rem;
        color: #1f4788;
        font-weight: bold;
        text-align: center;
        margin-bottom: 1.2rem;
    }
    .section-header {
        font-size: 1.4rem;
        color: #2c5aa0;
        font-weight: bold;
        margin-top: 0.8rem;
        margin-bottom: 0.6rem;
    }
    .stButton > button {
        background-color: #1f4788;
        color: white;
        font-weight: bold;
        padding: 0.5rem 2rem;
        border-radius: 5px;
        border: none;
        width: 100%;
    }
    .stButton > button:hover {
        background-color: #2c5aa0;
    }
    /* Reduce spacing between elements */
    .stTextInput > div > div > input {
        margin-bottom: 0;
    }
    .stSelectbox > div > div {
        margin-bottom: 0;
    }
    .stTextArea > div > div > textarea {
        margin-bottom: 0;
    }
    div[data-testid="stVerticalBlock"] > div {
        gap: 0.8rem;
    }
    .element-container {
        margin-bottom: 0.5rem;
    }
    h3 {
        margin-top: 1rem !important;
        margin-bottom: 0.5rem !important;
    }
    h4 {
        margin-top: 0.8rem !important;
        margin-bottom: 0.4rem !important;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'report_data' not in st.session_state:
    st.session_state.report_data = {}
if 'report_generated' not in st.session_state:
    st.session_state.report_generated = False
if 'technician_signature' not in st.session_state:
    st.session_state.technician_signature = None
if 'equipment_list' not in st.session_state:
    st.session_state.equipment_list = []
if 'form_data' not in st.session_state:
    st.session_state.form_data = {}


def render_checklist_item(equipment, item, equip_idx, prefix=""):
    """Recursively render a checklist item with all its conditional logic"""
    item_key = prefix + item['id']
    
    # Ensure inspection_data exists
    if 'inspection_data' not in equipment:
        equipment['inspection_data'] = {}
    
    if item_key not in equipment['inspection_data']:
        equipment['inspection_data'][item_key] = {}
    
    item_data = equipment['inspection_data'][item_key]
    
    # Render the question
    question = item['question']
    question_type = item['type']
    
    if question_type == 'yes_no':
        options = ['', 'Yes', 'No']
        widget_key = f"q_{item_key}_{equip_idx}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
            
        answer = st.selectbox(
            question,
            options=options,
            key=widget_key
        )
        # Update from session state
        st.session_state.equipment_list[equip_idx]['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
    elif question_type == 'yes_no_na':
        options = ['', 'Yes', 'No', 'N/A']
        widget_key = f"q_{item_key}_{equip_idx}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
            
        answer = st.selectbox(
            question,
            options=options,
            key=widget_key
        )
        # Update from session state
        st.session_state.equipment_list[equip_idx]['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
    elif question_type == 'text':
        widget_key = f"q_{item_key}_{equip_idx}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', '')
            
        answer = st.text_input(
            question,
            key=widget_key
        )
        # Update from session state
        st.session_state.equipment_list[equip_idx]['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
        
    elif question_type == 'number':
        widget_key = f"q_{item_key}_{equip_idx}"
        
        # Initialize session state if not exists
        if widget_key not in st.session_state:
            st.session_state[widget_key] = item_data.get('answer', 0)
            
        answer = st.number_input(
            question,
            min_value=0,
            step=1,
            key=widget_key
        )
        # Update from session state
        st.session_state.equipment_list[equip_idx]['inspection_data'][item_key]['answer'] = st.session_state[widget_key]
    else:
        answer = None
    
    # Get answer from updated data for conditional logic
    answer = st.session_state.equipment_list[equip_idx]['inspection_data'][item_key].get('answer', '')
    
    # Handle conditional logic based on answer
    if answer and 'conditions' in item:
        # Convert answer to lowercase for matching
        answer_key = answer.lower()
        condition = item['conditions'].get(answer_key)
        if condition:
            # Show a visual indicator that this answer triggered conditions
            st.markdown(f"<small style='color: #1f4788; margin: 0;'>→ Based on your answer '{answer}':</small>", unsafe_allow_html=True)
            
            # Handle photo requirement
            if condition.get('photo'):
                photo_key = f"photo_{item_key}"
                uploaded_files = st.file_uploader(
                    f"📷 Upload photo(s) for: {question}",
                    type=['png', 'jpg', 'jpeg'],
                    key=f"photo_{item_key}_{equip_idx}",
                    accept_multiple_files=True
                )
                if uploaded_files:
                    if 'photos' not in equipment:
                        equipment['photos'] = {}
                    # Store multiple photos with numbered keys
                    for i, uploaded_file in enumerate(uploaded_files):
                        if len(uploaded_files) == 1:
                            equipment['photos'][photo_key] = uploaded_file
                        else:
                            equipment['photos'][f"{photo_key}_{i+1}"] = uploaded_file
                    st.success(f"✅ {len(uploaded_files)} photo(s) uploaded")
            
            # Handle comment requirement
            if condition.get('comment'):
                comment_key = f"comment_{item_key}_{equip_idx}"
                
                # Initialize session state if not exists
                if comment_key not in st.session_state:
                    st.session_state[comment_key] = item_data.get('comment', '')
                    
                comment = st.text_area(
                    f"Please provide details for: {question}",
                    key=comment_key,
                    height=100
                )
                st.session_state.equipment_list[equip_idx]['inspection_data'][item_key]['comment'] = st.session_state[comment_key]
            
            # Handle action instruction
            if condition.get('action'):
                st.warning(condition['action'])
            
            # Handle follow-up questions
            if condition.get('follow_up'):
                # Add indentation for follow-up questions
                with st.container():
                    # Create a container with padding for visual hierarchy
                    _, col2 = st.columns([0.05, 0.95])
                    with col2:
                        for follow_up_item in condition['follow_up']:
                            render_checklist_item(equipment, follow_up_item, equip_idx, prefix)


def find_question_text(equipment_type, item_key):
    """Find the actual question text for a given item key"""
    def search_checklist(checklist_items, target_id):
        for item in checklist_items:
            if item['id'] == target_id:
                return item['question']
            # Check in follow-up questions
            if 'conditions' in item:
                for condition in item['conditions'].values():
                    if 'follow_up' in condition:
                        result = search_checklist(condition['follow_up'], target_id)
                        if result:
                            return result
        return None
    
    # Check if this is a Marvel question
    if item_key.startswith('marvel_'):
        # Remove marvel_ prefix and search in MARVEL checklist
        marvel_key = item_key[7:]  # Remove 'marvel_' prefix
        question = search_checklist(EQUIPMENT_TYPES.get('MARVEL', {}).get('checklist', []), marvel_key)
        if question:
            return question
    
    # Get the last part of the key which is the actual question ID
    parts = item_key.split('_')
    for i in range(len(parts)):
        # Try different combinations starting from the end
        for j in range(len(parts), i, -1):
            test_id = '_'.join(parts[i:j])
            question = search_checklist(EQUIPMENT_TYPES[equipment_type]['checklist'], test_id)
            if question:
                return question
    
    # Fallback to formatted key
    return item_key.replace('_', ' ').title()


def get_equipment_summary():
    """Get a summary of all equipment inspections"""
    summary = []
    
    for equipment in st.session_state.equipment_list:
        if equipment.get('type'):  # Only include equipment with a selected type
            equip_summary = {
                'type': equipment['type'],
                'type_name': EQUIPMENT_TYPES[equipment['type']]['name'],
                'with_marvel': equipment.get('with_marvel', False),
                'location': equipment.get('location', ''),
                'yes_responses': [],
                'no_responses': [],
                'photos_count': len(equipment.get('photos', {})),
                'inspection_data': equipment.get('inspection_data', {}),
                'photos': equipment.get('photos', {}),
                'yes_photos': {},
                'no_photos': {}
            }
            
            # Define items to exclude from No responses (these are not issues)
            exclude_from_no = ['final_remarks', 'lights_ballast']
            
            # Separate Yes and No responses
            for key, data in equipment.get('inspection_data', {}).items():
                if isinstance(data, dict):
                    answer = data.get('answer', '')
                    # Extract the base key - get the actual question ID
                    # Keys might be like "lights_ballast" or "lights_ballast_ballast_issue"
                    # We want to check if any part of the key is in the exclude list
                    key_parts = key.split('_')
                    should_exclude = False
                    for i in range(len(key_parts)):
                        check_key = '_'.join(key_parts[i:])
                        if check_key in exclude_from_no:
                            should_exclude = True
                            break
                    
                    # Get the actual question text
                    question_text = find_question_text(equipment['type'], key)
                    
                    if answer == 'Yes':
                        equip_summary['yes_responses'].append({
                            'item': key,
                            'question': question_text,
                            'answer': answer,
                            'comment': data.get('comment', '')
                        })
                        # Collect photos for this Yes response
                        photo_key = f"photo_{key}"
                        for pk, pv in equipment.get('photos', {}).items():
                            if pk.startswith(photo_key):
                                equip_summary['yes_photos'][pk] = pv
                    elif answer == 'No':
                        # Only add to no_responses if it's not in the exclude list
                        if not should_exclude:
                            equip_summary['no_responses'].append({
                                'item': key,
                                'question': question_text,
                                'answer': answer,
                                'comment': data.get('comment', '')
                            })
                            # Collect photos for this No response
                            photo_key = f"photo_{key}"
                            for pk, pv in equipment.get('photos', {}).items():
                                if pk.startswith(photo_key):
                                    equip_summary['no_photos'][pk] = pv
            
            # For backward compatibility, keep issues_found as no_responses
            equip_summary['issues_found'] = equip_summary['no_responses']
            
            summary.append(equip_summary)
    
    return summary


def create_technical_report(data):
    """Generate a Professional Technical Report Word document"""
    # Load the template document
    template_path = os.path.join(os.path.dirname(__file__), 'Templates', 'Report Letter Head.docx')
    
    # Check if template exists, otherwise create new document
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        # Fallback to creating new document if template not found
        doc = Document()
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.header_distance = Inches(0.5)
            section.footer_distance = Inches(0.5)
    
    # Add some space after letterhead
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add title manually to avoid underline
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run('TECHNICAL REPORT')
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(31, 71, 136)  # Halton Blue
    title_run.font.bold = True
    
    # Add report reference and date
    ref_para = doc.add_paragraph()
    ref_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref_para.add_run(f"Report Date: {data.get('date', datetime.now().strftime('%B %d, %Y'))}")
    ref_run.font.size = Pt(11)
    ref_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Add minimal spacing
    doc.add_paragraph()
    
    # GENERAL INFORMATION SECTION
    general_heading = doc.add_heading('1. GENERAL INFORMATION', level=1)
    style_heading(general_heading, level=1)
    
    # Create professional info table
    general_info = [
        ("Customer Name", data.get('customer_name', '')),
        ("Project Name", data.get('project_name', '')),
        ("Contact Person", data.get('contact_person', '')),
        ("Location", data.get('outlet_location', '')),
        ("Contact Number", data.get('contact_number', '')),
        ("Visit Type", data.get('visit_type', '')),
        ("Visit Classification", data.get('visit_class', ''))
    ]
    
    create_info_table(doc, general_info)
    doc.add_paragraph()  # Add spacing
    
    # EQUIPMENT INSPECTION SECTION
    equipment_heading = doc.add_heading('2. EQUIPMENT INSPECTION DETAILS', level=1)
    style_heading(equipment_heading, level=1)
    
    equipment_summary = data.get('equipment_inspection', [])
    
    if equipment_summary:
        for idx, equip in enumerate(equipment_summary):
            # Equipment header
            marvel_status = " (With Marvel)" if equip.get('with_marvel', False) else ""
            equip_title = doc.add_heading(f"{equip['type_name']}{marvel_status}", level=2)
            style_heading(equip_title, level=2)
            
            # Equipment info
            equip_info_para = doc.add_paragraph()
            equip_info_para.add_run(f"Location: {equip['location']}\n").font.size = Pt(11)
            equip_info_para.add_run(f"Total Photos Taken: {equip['photos_count']}\n").font.size = Pt(11)
            
            # YES RESPONSES SECTION
            if equip.get('yes_responses'):
                doc.add_paragraph()
                yes_heading = doc.add_paragraph()
                yes_run = yes_heading.add_run("Positive Findings:")
                yes_run.bold = True
                yes_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
                
                # Create table for yes responses
                yes_table_data = []
                for yes_item in equip['yes_responses']:
                    question_text = yes_item.get('question', yes_item['item'].replace('_', ' ').title())
                    answer_text = "YES"
                    if yes_item['comment']:
                        answer_text += f"\n{yes_item['comment']}"
                    yes_table_data.append((question_text, answer_text))
                
                if yes_table_data:
                    create_info_table(doc, yes_table_data, col_widths=[4, 2.5])
                
                # Add Yes photos if available
                if equip.get('yes_photos'):
                    doc.add_paragraph()
                    yes_photos_para = doc.add_paragraph()
                    yes_photos_para.add_run("Photos - Positive Findings:\n").bold = True
                    
                    # Group photos in pairs for side-by-side display
                    photo_items = list(equip['yes_photos'].items())
                    for i in range(0, len(photo_items), 2):
                        # Create a table for side-by-side photos
                        photo_table = doc.add_table(rows=1, cols=2)
                        photo_table.autofit = False
                        
                        # First photo
                        photo_key, photo_file = photo_items[i]
                        cell1 = photo_table.cell(0, 0)
                        cell1_para = cell1.paragraphs[0]
                        cell1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Reset file position and add photo
                        photo_file.seek(0)
                        run1 = cell1_para.add_run()
                        run1.add_picture(photo_file, width=Inches(2.0))
                        
                        # Add caption
                        caption1 = cell1.add_paragraph()
                        caption1.add_run(photo_key.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                        caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Second photo (if exists)
                        if i + 1 < len(photo_items):
                            photo_key2, photo_file2 = photo_items[i + 1]
                            cell2 = photo_table.cell(0, 1)
                            cell2_para = cell2.paragraphs[0]
                            cell2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Reset file position and add photo
                            photo_file2.seek(0)
                            run2 = cell2_para.add_run()
                            run2.add_picture(photo_file2, width=Inches(2.0))
                            
                            # Add caption
                            caption2 = cell2.add_paragraph()
                            caption2.add_run(photo_key2.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                            caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing after photos
                        doc.add_paragraph()
            
            # NO RESPONSES SECTION (Issues)
            if equip.get('no_responses'):
                doc.add_paragraph()
                no_heading = doc.add_paragraph()
                no_run = no_heading.add_run("Issues Identified:")
                no_run.bold = True
                no_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                
                # Create table for no responses
                no_table_data = []
                for no_item in equip['no_responses']:
                    question_text = no_item.get('question', no_item['item'].replace('_', ' ').title())
                    answer_text = "NO"
                    if no_item['comment']:
                        answer_text += f"\n{no_item['comment']}"
                    no_table_data.append((question_text, answer_text))
                
                if no_table_data:
                    create_info_table(doc, no_table_data, col_widths=[4, 2.5])
                
                # Add No photos if available
                if equip.get('no_photos'):
                    doc.add_paragraph()
                    no_photos_para = doc.add_paragraph()
                    no_photos_para.add_run("Photos - Issues:\n").bold = True
                    
                    # Group photos in pairs for side-by-side display
                    photo_items = list(equip['no_photos'].items())
                    for i in range(0, len(photo_items), 2):
                        # Create a table for side-by-side photos
                        photo_table = doc.add_table(rows=1, cols=2)
                        photo_table.autofit = False
                        
                        # First photo
                        photo_key, photo_file = photo_items[i]
                        cell1 = photo_table.cell(0, 0)
                        cell1_para = cell1.paragraphs[0]
                        cell1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Reset file position and add photo
                        photo_file.seek(0)
                        run1 = cell1_para.add_run()
                        run1.add_picture(photo_file, width=Inches(2.0))
                        
                        # Add caption
                        caption1 = cell1.add_paragraph()
                        caption1.add_run(photo_key.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                        caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Second photo (if exists)
                        if i + 1 < len(photo_items):
                            photo_key2, photo_file2 = photo_items[i + 1]
                            cell2 = photo_table.cell(0, 1)
                            cell2_para = cell2.paragraphs[0]
                            cell2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Reset file position and add photo
                            photo_file2.seek(0)
                            run2 = cell2_para.add_run()
                            run2.add_picture(photo_file2, width=Inches(2.0))
                            
                            # Add caption
                            caption2 = cell2.add_paragraph()
                            caption2.add_run(photo_key2.replace('photo_', '').replace('_', ' ').title()).font.size = Pt(9)
                            caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add spacing after photos
                        doc.add_paragraph()
            
            # If no issues found at all
            if not equip.get('no_responses'):
                doc.add_paragraph()
                para = doc.add_paragraph()
                no_issues_run = para.add_run("No issues identified during inspection.")
                no_issues_run.font.size = Pt(11)
                no_issues_run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            
            # Add spacing between equipment
            if idx < len(equipment_summary) - 1:
                doc.add_paragraph()
    else:
        para = doc.add_paragraph()
        para.add_run("No equipment inspection data available.").font.size = Pt(11)
    
    # WORK PERFORMED SECTION
    work_heading = doc.add_heading('3. JOB DETAILS', level=1)
    style_heading(work_heading, level=1)
    
    work_para = doc.add_paragraph()
    work_text = work_para.add_run(data.get('work_performed', ''))
    work_text.font.size = Pt(11)
    work_para.paragraph_format.line_spacing = 1.5
    work_para.paragraph_format.space_after = Pt(12)
    
    # RECOMMENDATIONS SECTION
    if data.get('recommendations'):
        rec_heading = doc.add_heading('4. RECOMMENDATIONS', level=1)
        style_heading(rec_heading, level=1)
        
        rec_para = doc.add_paragraph()
        rec_text = rec_para.add_run(data.get('recommendations', ''))
        rec_text.font.size = Pt(11)
        rec_para.paragraph_format.line_spacing = 1.5
        rec_para.paragraph_format.space_after = Pt(12)
    
    # Add page break before signatures
    doc.add_page_break()
    
    # SIGNATURE SECTION
    sig_heading = doc.add_heading('ACKNOWLEDGMENT AND SIGNATURES', level=1)
    style_heading(sig_heading, level=1)
    
    # Add acknowledgment text
    ack_para = doc.add_paragraph()
    ack_text = ack_para.add_run(
        "The undersigned acknowledge that the service described in this report has been "
        "completed satisfactorily and in accordance with the agreed specifications."
    )
    ack_text.font.size = Pt(10)
    ack_text.font.italic = True
    ack_para.paragraph_format.space_after = Pt(24)
    
    # Create signature table
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Configure signature columns
    for col in sig_table.columns:
        for cell in col.cells:
            cell.width = Inches(3)
    
    # Technician signature
    sig_table.cell(0, 0).text = "Service Technician:"
    
    # Add signature image if available
    if data.get('technician_signature'):
        sig_cell = sig_table.cell(1, 0)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sig_para.add_run()
        run.add_picture(data.get('technician_signature'), width=Inches(1.5))
    else:
        sig_table.cell(1, 0).text = "_" * 35
    
    sig_table.cell(2, 0).text = data.get('technician_name', '')
    sig_table.cell(3, 0).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Customer signature
    sig_table.cell(0, 1).text = "Customer Representative:"
    
    # Add customer signature image if available
    if data.get('customer_signature'):
        sig_cell = sig_table.cell(1, 1)
        sig_para = sig_cell.paragraphs[0]
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sig_para.add_run()
        run.add_picture(data.get('customer_signature'), width=Inches(1.5))
    else:
        sig_table.cell(1, 1).text = "_" * 35
    
    sig_table.cell(2, 1).text = data.get('customer_signatory', data.get('customer_name', ''))
    sig_table.cell(3, 1).text = f"Date: {datetime.now().strftime('%B %d, %Y')}"
    
    # Style signature table
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
            set_cell_margins(cell, top=0.1, bottom=0.1)
    
    # Make labels bold
    sig_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    sig_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
    # Add final note
    doc.add_paragraph()
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_text = note_para.add_run(
        "This report is confidential and proprietary to Halton Company Saudi Arabia Ltd.\n"
        "For service inquiries, please contact our Service Department."
    )
    note_text.font.size = Pt(9)
    note_text.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    return doc_bytes

def main():
    # Header
    st.markdown('<h1 class="main-header">Halton KSA Service Reports</h1>', unsafe_allow_html=True)
    
    # Sidebar for report type selection
    with st.sidebar:
        st.markdown("### Report Type Selection")
        report_type = st.selectbox(
            "Select Report Type",
            ["Technical Report", "Testing and Commissioning Report (Coming Soon)", "General Service Report (Coming Soon)"],
            index=0
        )
        
        if report_type != "Technical Report":
            st.info("This report type will be available in future versions.")
            st.stop()
    
    # Main form for Technical Report
    st.markdown('<h2 class="section-header">Technical Report Form</h2>', unsafe_allow_html=True)
    
    # General Information Section (outside form)
    st.markdown("### General Information")
    col1, col2 = st.columns(2)
    
    with col1:
        customer_name = st.text_input("Customer's Name*", placeholder="e.g., SELA Company", key="customer_name")
        project_name = st.text_input("Project Name*", placeholder="e.g., Stella kitchen hoods", key="project_name")
        contact_person = st.text_input("Contact Person*", placeholder="e.g., Sultan Alofi", key="contact_person")
        outlet_location = st.text_input("Outlet/Location*", placeholder="e.g., Via - Riyadh", key="outlet_location")
    
    with col2:
        contact_number = st.text_input("Contact #*", placeholder="e.g., +966 55 558 5449", key="contact_number")
        visit_type = st.selectbox(
            "Visit Type*",
            ["", "Servicing/Preventive Maintenance", "AMC (Contract)", "Emergency Service", "Installation", "Commissioning"],
            key="visit_type"
        )
        
        visit_class = st.selectbox(
            "Visit Class*",
            ["To be invoiced (Chargeable)", "Warranty", "Complaint", "Scheduled Maintenance", "Emergency Service"],
            key="visit_class"
        )
        date = st.date_input("Report Date", value=datetime.now(), key="report_date")
    
    # Equipment Inspection Section (outside form for real-time updates)
    st.markdown("### Equipment Inspection")
    
    # Number of equipment to inspect
    def update_equipment_count():
        # This callback ensures equipment list updates properly
        pass
    
    num_equipment = st.number_input(
        "Number of Equipment to Inspect",
        min_value=1,
        max_value=10,
        value=len(st.session_state.equipment_list) if st.session_state.equipment_list else 1,
        step=1,
        key="num_equipment",
        on_change=update_equipment_count
    )
    
    # Initialize equipment list if needed
    if len(st.session_state.equipment_list) != num_equipment:
        # Adjust equipment list size
        if len(st.session_state.equipment_list) < num_equipment:
            # Add new equipment
            for i in range(len(st.session_state.equipment_list), num_equipment):
                equipment_id = f"equipment_{i}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                st.session_state.equipment_list.append({
                    'id': equipment_id,
                    'type': '',
                    'with_marvel': False,
                    'location': '',
                    'inspection_data': {},
                    'photos': {}
                })
        else:
            # Remove equipment
            st.session_state.equipment_list = st.session_state.equipment_list[:num_equipment]
    
    # Display equipment forms
    for idx, equipment in enumerate(st.session_state.equipment_list):
        with st.expander(f"Equipment #{idx + 1}", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                # Equipment type selection
                equip_type_key = f"equip_type_{idx}"
                # Initialize if not exists
                if equip_type_key not in st.session_state:
                    st.session_state[equip_type_key] = equipment.get('type', '')
                    
                st.selectbox(
                    "Equipment Type*",
                    options=[''] + list(EQUIPMENT_TYPES.keys()),
                    format_func=lambda x: EQUIPMENT_TYPES[x]["name"] if x else "Select equipment type",
                    key=equip_type_key
                )
                # Update the equipment data from session state
                equipment['type'] = st.session_state[equip_type_key]
                
                marvel_key = f"with_marvel_{idx}"
                # Initialize if not exists
                if marvel_key not in st.session_state:
                    st.session_state[marvel_key] = equipment.get('with_marvel', False)
                    
                with_marvel = st.checkbox(
                    "With Marvel System",
                    key=marvel_key
                )
                # Update the equipment data from session state
                equipment['with_marvel'] = st.session_state[marvel_key]
            
            with col2:
                location_key = f"location_{idx}"
                # Initialize if not exists
                if location_key not in st.session_state:
                    st.session_state[location_key] = equipment.get('location', '')
                    
                st.text_input(
                    "Location*",
                    key=location_key
                )
                # Update the equipment data from session state
                equipment['location'] = st.session_state[location_key]
            
            # If equipment type is selected, show checklist
            if equipment['type']:
                st.markdown("#### Inspection Checklist")
                equipment_config = EQUIPMENT_TYPES[equipment['type']]
                
                # Render checklist items with full conditional logic
                for i, item in enumerate(equipment_config['checklist']):
                    if i > 0:
                        st.markdown("<hr style='margin: 0.5rem 0;'>", unsafe_allow_html=True)  # Thinner separator
                    render_checklist_item(equipment, item, idx)
                
                # If "With Marvel" is checked, add Marvel checklist
                if equipment.get('with_marvel', False):
                    st.markdown("#### Marvel System Checklist")
                    marvel_config = EQUIPMENT_TYPES.get('MARVEL', {})
                    if marvel_config and 'checklist' in marvel_config:
                        for i, item in enumerate(marvel_config['checklist']):
                            if i > 0:
                                st.markdown("<hr style='margin: 0.5rem 0;'>", unsafe_allow_html=True)
                            # Add prefix to distinguish Marvel questions
                            render_checklist_item(equipment, item, idx, prefix="marvel_")
    
    # Continue with the rest of the form
    with st.form("technical_report_form"):
        # Work Performed Section
        st.markdown("### Work Performed")
        work_performed = st.text_area(
            "Describe Work Performed*",
            placeholder="Detail all maintenance, repairs, or services completed...",
            height=120
        )
        
        
        # Recommendations Section
        st.markdown("### Recommendations")
        recommendations = st.text_area(
            "Recommendations",
            placeholder="Suggest any follow-up actions, parts needed, or future maintenance...",
            height=80
        )
        
        # Technician Information Section
        st.markdown("### Technician Information")
        col3, col4 = st.columns(2)
        
        with col3:
            technician_name = st.text_input("Technician Name*", placeholder="Enter your full name")
            technician_id = st.text_input("Technician ID*", placeholder="Enter your employee ID")
        
        with col4:
            service_date = st.date_input("Service Date", value=datetime.now())
        
        # Signature Section
        st.markdown("### Technician Signature")
        st.markdown("Please draw your signature below using your mouse or touchscreen")
        
        # Create signature canvas
        col1, col2 = st.columns([4, 1])
        
        with col1:
            canvas_result = st_canvas(
                fill_color="rgba(255, 255, 255, 0)",  # Transparent fill
                stroke_width=3,
                stroke_color="#000000",
                background_color="#FFFFFF",
                background_image=None,
                update_streamlit=True,
                height=120,
                width=450,
                drawing_mode="freedraw",
                point_display_radius=0,
                display_toolbar=True,
                key="signature_canvas",
            )
        
        with col2:
            st.markdown("### ")  # Add spacing
            st.info("Use the trash icon in the canvas toolbar to clear")
        
        # Check if signature is drawn
        if canvas_result.image_data is not None:
            # Check if canvas has any drawing (non-transparent pixels)
            if np.any(canvas_result.image_data[:,:,3] > 0):
                st.success("✅ Technician signature captured")
        
        # Customer Signature Section
        st.markdown("### Customer Signature")
        col3, col4 = st.columns(2)
        
        with col3:
            # Customer name field - default to customer name from general info
            st.text_input(
                "Customer Representative Name",
                value=st.session_state.get('customer_name', ''),
                placeholder="Enter customer representative name",
                key="customer_signatory"
            )
        
        with col4:
            st.markdown("#### ")  # Spacing
            
        st.markdown("Please have the customer draw their signature below")
        
        # Create customer signature canvas
        col5, col6 = st.columns([4, 1])
        
        with col5:
            customer_canvas_result = st_canvas(
                fill_color="rgba(255, 255, 255, 0)",  # Transparent fill
                stroke_width=3,
                stroke_color="#000000",
                background_color="#FFFFFF",
                background_image=None,
                update_streamlit=True,
                height=120,
                width=450,
                drawing_mode="freedraw",
                point_display_radius=0,
                display_toolbar=True,
                key="customer_signature_canvas",
            )
        
        with col6:
            st.markdown("### ")  # Add spacing
            st.info("Use the trash icon in the canvas toolbar to clear")
        
        # Check if customer signature is drawn
        if customer_canvas_result.image_data is not None:
            # Check if canvas has any drawing (non-transparent pixels)
            if np.any(customer_canvas_result.image_data[:,:,3] > 0):
                st.success("✅ Customer signature captured")
        
        # Submit button
        submitted = st.form_submit_button("Generate Report", type="primary")
        
        if submitted:
            # Get values from session state
            customer_name = st.session_state.get('customer_name', '')
            project_name = st.session_state.get('project_name', '')
            contact_person = st.session_state.get('contact_person', '')
            outlet_location = st.session_state.get('outlet_location', '')
            contact_number = st.session_state.get('contact_number', '')
            visit_type = st.session_state.get('visit_type', '')
            visit_class = st.session_state.get('visit_class', '')
            date = st.session_state.get('report_date', datetime.now())
            
            # Validation
            required_fields = {
                "Customer's Name": customer_name,
                "Project Name": project_name,
                "Contact Person": contact_person,
                "Outlet/Location": outlet_location,
                "Contact Number": contact_number,
                "Visit Type": visit_type,
                "Work Performed": work_performed,
                "Technician Name": technician_name,
                "Technician ID": technician_id
            }
            
            missing_fields = [field for field, value in required_fields.items() if not value]
            
            # Validate equipment data
            equipment_errors = []
            for idx, equipment in enumerate(st.session_state.equipment_list):
                equip_name = f"Equipment #{idx + 1}"
                if not equipment.get('type'):
                    equipment_errors.append(f"{equip_name}: Equipment type is required")
                elif not equipment.get('location'):
                    equipment_errors.append(f"{equip_name}: Location is required")
            
            # Show reminders for missing fields but don't block submission
            if missing_fields or equipment_errors:
                st.warning("📋 **Reminder:** The following fields are recommended but not required:")
                if missing_fields:
                    st.write("**General Information:**")
                    for field in missing_fields:
                        st.write(f"• {field}")
                if equipment_errors:
                    st.write("**Equipment Information:**")
                    for error in equipment_errors:
                        st.write(f"• {error}")
                st.info("The report will be generated with the available information.")
            
            # Always proceed with report generation
            # Process signature from canvas
            signature_img = None
            if canvas_result.image_data is not None and np.any(canvas_result.image_data[:,:,3] > 0):
                # Convert canvas to image
                sig_image = Image.fromarray(canvas_result.image_data.astype('uint8'), 'RGBA')
                
                # Create white background
                white_bg = Image.new('RGB', sig_image.size, 'white')
                white_bg.paste(sig_image, mask=sig_image.split()[3])
                
                # Find the bounding box of the signature
                bbox = white_bg.getbbox()
                if bbox:
                    # Crop to signature
                    cropped = white_bg.crop(bbox)
                    
                    # Resize if too large
                    max_width, max_height = 200, 80
                    cropped.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    
                    # Convert to bytes
                    img_bytes = io.BytesIO()
                    cropped.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    signature_img = img_bytes
            
            # Process customer signature from canvas
            customer_signature_img = None
            if customer_canvas_result.image_data is not None and np.any(customer_canvas_result.image_data[:,:,3] > 0):
                # Convert canvas to image
                cust_sig_image = Image.fromarray(customer_canvas_result.image_data.astype('uint8'), 'RGBA')
                
                # Create white background
                white_bg = Image.new('RGB', cust_sig_image.size, 'white')
                white_bg.paste(cust_sig_image, mask=cust_sig_image.split()[3])
                
                # Find the bounding box of the signature
                bbox = white_bg.getbbox()
                if bbox:
                    # Crop to signature
                    cropped = white_bg.crop(bbox)
                    
                    # Resize if too large
                    max_width, max_height = 200, 80
                    cropped.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
                    
                    # Convert to bytes
                    img_bytes = io.BytesIO()
                    cropped.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    customer_signature_img = img_bytes
            
            # Collect all data
            report_data = {
                    'customer_name': customer_name,
                    'project_name': project_name,
                    'contact_person': contact_person,
                    'outlet_location': outlet_location,
                    'contact_number': contact_number,
                    'visit_type': visit_type,
                    'visit_class': visit_class,
                    'date': date.strftime('%Y-%m-%d'),
                    'equipment_inspection': get_equipment_summary(),
                    'equipment_list': st.session_state.equipment_list,
                    'work_performed': work_performed,
                    'recommendations': recommendations,
                    'technician_name': technician_name,
                    'technician_id': technician_id,
                    'service_date': service_date.strftime('%Y-%m-%d'),
                    'technician_signature': signature_img,
                    'customer_signatory': st.session_state.get('customer_signatory', customer_name),
                    'customer_signature': customer_signature_img
            }
            
            # Store data in session state for download outside form
            st.session_state.report_data = report_data
            st.session_state.report_generated = True
            st.session_state.saved_customer_name = customer_name
            st.session_state.saved_report_date = date
    
    # Handle report download outside of form
    if st.session_state.get('report_generated', False):
        try:
            # Generate the report
            doc_bytes = create_technical_report(st.session_state.report_data)
            
            # Create filename
            customer_name = st.session_state.saved_customer_name
            date = st.session_state.saved_report_date
            filename = f"Technical_Report_{customer_name.replace(' ', '_')}_{date.strftime('%Y%m%d')}.docx"
            
            # Success message
            st.success("✅ Report generated successfully!")
            
            # Download button (outside form)
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                st.download_button(
                    label="📥 Download Report",
                    data=doc_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            # Option to generate another report
            if st.button("Generate Another Report", type="secondary"):
                # Clear all session state data for a fresh start
                st.session_state.report_generated = False
                st.session_state.equipment_list = []
                st.session_state.report_data = {}
                # Clear all widget keys
                keys_to_clear = []
                for key in st.session_state.keys():
                    if (key.startswith('q_') or key.startswith('comment_') or 
                        key.startswith('equip_type_') or key.startswith('with_marvel_') or 
                        key.startswith('location_') or key.startswith('photo_') or
                        key == 'customer_signatory' or key == 'customer_signature_canvas' or
                        key == 'signature_canvas'):
                        keys_to_clear.append(key)
                for key in keys_to_clear:
                    del st.session_state[key]
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ Error generating report: {str(e)}")
            if st.button("Try Again"):
                st.session_state.report_generated = False
                st.rerun()

if __name__ == "__main__":
    main()